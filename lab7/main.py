import sys
from PyQt6.QtWidgets import *
from docx import Document
from geometry_pkg import Parallelepiped, Tetrahedron, Sphere

MATERIALS = {"Алюминий": 2.70, "Сталь": 7.85, "Пластик": 1.38}

class GeometryApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Расчёт геометрических тел")
        self.setGeometry(100, 100, 500, 450)
        self.current_body = None
        self.inputs = []  # список полей ввода
        self.setup_ui()
    
    def setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        
        # Выбор тела
        self.body_combo = QComboBox()
        self.body_combo.addItems(["Параллелепипед", "Тетраэдр", "Шар"])
        self.body_combo.currentTextChanged.connect(self.update_params)
        layout.addWidget(self.body_combo)
        
        # Контейнер для параметров (будет очищаться)
        self.params_container = QWidget()
        self.params_layout = QVBoxLayout(self.params_container)
        layout.addWidget(self.params_container)
        
        # Материал
        layout.addWidget(QLabel("Материал:"))
        self.material_combo = QComboBox()
        self.material_combo.addItems(list(MATERIALS.keys()))
        layout.addWidget(self.material_combo)
        
        # Кнопка расчёта
        btn = QPushButton("Рассчитать")
        btn.clicked.connect(self.calculate)
        layout.addWidget(btn)
        
        # Поле результата
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        layout.addWidget(self.result_text)
        
        # Кнопка сохранения
        save_btn = QPushButton("Сохранить в .docx")
        save_btn.clicked.connect(self.save_report)
        layout.addWidget(save_btn)
        
        # Инициализируем поля для тела по умолчанию
        self.update_params(self.body_combo.currentText())
    
    def update_params(self, body_type):
        # Очищаем контейнер с параметрами
        while self.params_layout.count():
            item = self.params_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                # удаляем вложенные layout (наши hbox)
                while item.layout().count():
                    subitem = item.layout().takeAt(0)
                    if subitem.widget():
                        subitem.widget().deleteLater()
                item.layout().deleteLater()
        
        self.inputs.clear()
        
        # Выбираем подписи в зависимости от тела
        if body_type == "Параллелепипед":
            labels = ["Длина (см):", "Ширина (см):", "Высота (см):"]
        elif body_type == "Тетраэдр":
            labels = ["Ребро (см):"]
        else:  # Шар
            labels = ["Радиус (см):"]
        
        # Создаём новые поля
        for lbl in labels:
            hbox = QHBoxLayout()
            hbox.addWidget(QLabel(lbl))
            inp = QLineEdit()
            hbox.addWidget(inp)
            self.params_layout.addLayout(hbox)
            self.inputs.append(inp)
    
    def get_params(self):
        try:
            return [float(inp.text()) for inp in self.inputs]
        except ValueError:
            QMessageBox.critical(self, "Ошибка", "Введите числа")
            return None
    
    def calculate(self):
        params = self.get_params()
        if not params:
            return
        
        body = self.body_combo.currentText()
        material = self.material_combo.currentText()
        density = MATERIALS[material]
        
        try:
            if body == "Параллелепипед":
                self.current_body = Parallelepiped(*params, material, density)
            elif body == "Тетраэдр":
                self.current_body = Tetrahedron(params[0], material, density)
            else:
                self.current_body = Sphere(params[0], material, density)
            
            print(str(self.current_body)) 
            print(repr(self.current_body)) 
            
            d = self.current_body.to_dict()
            self.result_text.setText(
                f"{self.current_body}\n"
                f"Плотность: {density} г/см³\n"
                f"Объём: {d['volume']:.2f} см³\n"
                f"Площадь: {d['surface_area']:.2f} см²\n"
                f"Масса: {d['mass']:.2f} г ({d['mass']/1000:.3f} кг)"
            )
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))
    
    def save_report(self):
        if not self.current_body:
            QMessageBox.warning(self, "Нет данных", "Сначала выполните расчёт")
            return
        
        path, _ = QFileDialog.getSaveFileName(self, "Сохранить", "", "Word (*.docx)")
        if not path:
            return
        
        d = self.current_body.to_dict()
        doc = Document()
        doc.add_paragraph(f"Тело: {d['body_type']}")
        doc.add_paragraph(f"Материал: {d['material']} (плотность {d['density']} г/см³)")
        doc.add_paragraph(f"Объём: {d['volume']:.2f} см³")
        doc.add_paragraph(f"Площадь поверхности: {d['surface_area']:.2f} см²")
        doc.add_paragraph(f"Масса: {d['mass']:.2f} г ({d['mass']/1000:.3f} кг)")
        doc.save(path)
        QMessageBox.information(self, "Готово", f"Сохранено: {path}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = GeometryApp()
    window.show()
    sys.exit(app.exec())