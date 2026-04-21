import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from geometry_pkg import parallelepiped, tetrahedron, sphere
from docx import Document

# материалы плотность в г/см3
MATERIALS = {
    "алюминий": 2.70,
    "сталь": 7.85,
    "пластик": 1.38,
}

class GeometryApp:
    def __init__(self, root):
        self.root = root
        self.root.title("расчет тел")
        self.root.geometry("500x450")

        self.body_var = tk.StringVar(value="параллелепипед")
        self.material_var = tk.StringVar(value="алюминий")
        self.entries = []

        # выбор тела
        tk.Label(root, text="тело:").pack()
        ttk.Combobox(root, textvariable=self.body_var, values=["параллелепипед", "тетраэдр", "шар"], state="readonly").pack()
        self.body_var.trace('w', self.update_params_frame)

        # рамка для параметров
        self.params_frame = tk.LabelFrame(root, text="параметры")
        self.params_frame.pack(pady=10, padx=10, fill='x')

        # материал
        tk.Label(root, text="материал:").pack()
        ttk.Combobox(root, textvariable=self.material_var, values=list(MATERIALS.keys()), state="readonly").pack()

        # кнопка расчёта
        tk.Button(root, text="рассчитать", command=self.calculate).pack(pady=10)

        # поле для результата
        self.result_text = tk.Text(root, height=8, width=60)
        self.result_text.pack(pady=10)

        # кнопка сохранения в документ
        tk.Button(root, text="сохранить в доки", command=self.save_to_docx).pack(pady=5)

        self.update_params_frame()
        self.last_result = None

    def update_params_frame(self, *args):
        # очищаем старые
        for widget in self.params_frame.winfo_children():
            widget.destroy()
        self.entries.clear()

        body = self.body_var.get()
        if body == "параллелепипед":
            labels = ["длина (см):", "ширина (см):", "высота (см):"]
        elif body == "тетраэдр":
            labels = ["ребро (см):"]
        else:  # Шар
            labels = ["радиус (см):"]

        for i, lbl in enumerate(labels):
            tk.Label(self.params_frame, text=lbl).grid(row=i, column=0, sticky='w')
            ent = tk.Entry(self.params_frame, width=15)
            ent.grid(row=i, column=1, padx=5, pady=2)
            self.entries.append(ent)

    def get_params(self):
        try:
            return [float(ent.get()) for ent in self.entries]
        except ValueError:
            messagebox.showerror("ошибка", "введите числа")
            return None

    def calculate(self):
        params = self.get_params()
        if params is None:
            return

        body = self.body_var.get()
        material = self.material_var.get()
        density = MATERIALS[material]

        # расчёт в зависимости от фигуры
        if body == "параллелепипед":
            a, b, c = params
            vol = parallelepiped.volume(a, b, c)
            area = parallelepiped.surface_area(a, b, c)
            mass_val = parallelepiped.mass(vol, density)
        elif body == "тетраэдр":
            a = params[0]
            vol = tetrahedron.volume(a)
            area = tetrahedron.surface_area(a)
            mass_val = tetrahedron.mass(vol, density)
        else:  # шарик
            r = params[0]
            vol = sphere.volume(r)
            area = sphere.surface_area(r)
            mass_val = sphere.mass(vol, density)

        result_str = (
            f"тело: {body}\n"
            f"параметры: {params} см\n"
            f"материал: {material} (плотность {density} г/см3)\n"
            f"объём: {vol:.2f} см3\n"
            f"площадь поверхности: {area:.2f} см2\n"
            f"масса: {mass_val:.2f} г  ({mass_val/1000:.3f} кг)"
        )

        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(tk.END, result_str)

        # сохрн
        self.last_result = {
            'body': body,
            'params': params,
            'material': material,
            'density': density,
            'volume': vol,
            'area': area,
            'mass': mass_val
        }

    def save_to_docx(self):
        if self.last_result is None:
            messagebox.showwarning("нужны данные", "нет расчета")
            return

        r = self.last_result
        path = filedialog.asksaveasfilename(defaultextension='.docx', filetypes=[("Word документ", "*.docx")])
        if not path:
            return

        doc = Document()
        doc.add_paragraph(f"тело: {r['body']}")
        doc.add_paragraph(f"параметры (см): {r['params']}")
        doc.add_paragraph(f"материал: {r['material']} (плотность {r['density']} г/см3)")
        doc.add_paragraph(f"объём: {r['volume']:.2f} см3")
        doc.add_paragraph(f"площадь поверхности: {r['area']:.2f} см2")
        doc.add_paragraph(f"масса: {r['mass']:.2f} г ({r['mass']/1000:.3f} кг)")

        doc.save(path)
        messagebox.showinfo("сохранено", f"файл сохранён: {path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = GeometryApp(root)
    root.mainloop()